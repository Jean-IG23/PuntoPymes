#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para generar Manual de Usuario de PuntoPymes en formato Word
Enfocado en las interacciones del usuario con el sistema
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level, color=None):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_title_page(doc):
    """Add title page"""
    # Title
    title = doc.add_heading('MANUAL DE USUARIO', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    subtitle = doc.add_paragraph('PuntoPymes v2.0')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 102, 204)
    
    # Description
    desc = doc.add_paragraph('Sistema Integral de Gestión de Recursos Humanos')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in desc.runs:
        run.font.size = Pt(14)
        run.font.italic = True
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Fecha
    fecha = doc.add_paragraph('Enero 2026')
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fecha.runs:
        run.font.size = Pt(12)
    
    doc.add_page_break()

def add_table_of_contents(doc):
    """Add table of contents"""
    add_heading_styled(doc, 'TABLA DE CONTENIDOS', level=1, color=RGBColor(0, 51, 102))
    
    toc_items = [
        '1. Introducción',
        '2. Acceso al Sistema',
        '3. Interfaz Principal',
        '4. Módulos Principales',
        '   4.1 Dashboard',
        '   4.2 Gestión de Empleados',
        '   4.3 Asistencia',
        '   4.4 Ausencias y Vacaciones',
        '   4.5 Objetivos y KPIs',
        '   4.6 Tareas',
        '   4.7 Nómina',
        '   4.8 Configuración',
        '5. Guía por Rol',
        '   5.1 Administrador',
        '   5.2 RRHH',
        '   5.3 Gerente',
        '   5.4 Empleado',
        '6. Operaciones Comunes',
        '7. Preguntas Frecuentes',
        '8. Solución de Problemas',
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()

def add_introduction(doc):
    """Add introduction section"""
    add_heading_styled(doc, '1. Introducción', level=1, color=RGBColor(0, 51, 102))
    
    doc.add_paragraph(
        'PuntoPymes es una plataforma integral de gestión de recursos humanos diseñada '
        'para empresas de cualquier tamaño. El sistema permite gestionar de forma eficiente '
        'la información de empleados, asistencia, vacaciones, objetivos y nóminas.'
    )
    
    add_heading_styled(doc, '¿Qué es PuntoPymes?', level=2)
    doc.add_paragraph(
        'PuntoPymes es una solución moderna que centraliza la gestión de recursos humanos. '
        'Con una interfaz intuitiva y funcionalidades avanzadas, permite a empresas de cualquier '
        'tamaño gestionar eficientemente a sus equipos.'
    )
    
    add_heading_styled(doc, 'Características Principales', level=2)
    features = [
        'Gestión centralizada de empleados y sucursales',
        'Sistema de asistencia en tiempo real',
        'Control de vacaciones y ausencias',
        'Definición y seguimiento de objetivos',
        'Sistema de tareas colaborativas',
        'Generación de nóminas',
        'Reportes y analíticas',
        'Control de acceso por roles y permisos'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()

def add_login_section(doc):
    """Add login section"""
    add_heading_styled(doc, '2. Acceso al Sistema', level=1, color=RGBColor(0, 51, 102))
    
    add_heading_styled(doc, 'Inicio de Sesión', level=2)
    doc.add_paragraph('Para acceder a PuntoPymes:')
    
    steps = [
        'Abre tu navegador web y ve a la URL de tu instancia de PuntoPymes',
        'Ingresa tu email o usuario en el campo "Usuario/Email"',
        'Ingresa tu contraseña en el campo "Contraseña"',
        'Haz clic en el botón "Ingresar"',
        'Si marcaste "Recuérdame", tu sesión se mantendrá activa'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    add_heading_styled(doc, '¿Olvidaste tu Contraseña?', level=2)
    doc.add_paragraph(
        '1. En la pantalla de login, haz clic en "¿Olvidaste tu contraseña?"'
    )
    doc.add_paragraph(
        '2. Ingresa tu email registrado'
    )
    doc.add_paragraph(
        '3. Recibirás un email con un link para reestablecerla'
    )
    doc.add_paragraph(
        '4. Sigue las instrucciones en el email'
    )
    
    doc.add_page_break()

def add_main_interface(doc):
    """Add main interface section"""
    add_heading_styled(doc, '3. Interfaz Principal', level=1, color=RGBColor(0, 51, 102))
    
    add_heading_styled(doc, 'Componentes de la Interfaz', level=2)
    doc.add_paragraph(
        'Una vez hayas iniciado sesión, verás la interfaz principal de PuntoPymes con los '
        'siguientes componentes:'
    )
    
    # Create interface components table
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Componente'
    header_cells[1].text = 'Descripción'
    
    # Set header background
    for cell in header_cells:
        set_cell_background(cell, '003366')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    components = [
        ('Barra Superior', 'Contiene el logo, nombre de la empresa y opciones de usuario'),
        ('Menú Lateral', 'Navegación entre módulos: Dashboard, Empleados, Asistencia, etc.'),
        ('Panel Principal', 'Área donde se muestra el contenido del módulo seleccionado'),
        ('Perfil de Usuario', 'Acceso a configuración personal y cerrar sesión'),
        ('Notificaciones', 'Alertas y mensajes del sistema'),
        ('Breadcrumbs', 'Ruta de navegación actual (ej: Inicio > Empleados > Editar)')
    ]
    
    for i, (component, description) in enumerate(components, 1):
        row = table.rows[i]
        row.cells[0].text = component
        row.cells[1].text = description
    
    doc.add_page_break()

def add_modules_section(doc):
    """Add modules section"""
    add_heading_styled(doc, '4. Módulos Principales', level=1, color=RGBColor(0, 51, 102))
    
    # Dashboard
    add_heading_styled(doc, '4.1 Dashboard', level=2)
    doc.add_paragraph(
        'El Dashboard es tu página de inicio. Aquí puedes ver un resumen de la información '
        'más importante según tu rol.'
    )
    doc.add_paragraph('Elementos del Dashboard:', style='List Bullet')
    doc.add_paragraph('KPIs y métricas clave del mes', style='List Bullet 2')
    doc.add_paragraph('Gráficos de asistencia', style='List Bullet 2')
    doc.add_paragraph('Alertas de vacaciones próximas', style='List Bullet 2')
    doc.add_paragraph('Tareas pendientes', style='List Bullet 2')
    doc.add_paragraph('Últimas actividades', style='List Bullet 2')
    
    # Empleados
    add_heading_styled(doc, '4.2 Gestión de Empleados', level=2)
    doc.add_paragraph(
        'En este módulo puedes gestionar toda la información de tus empleados.'
    )
    
    add_heading_styled(doc, 'Listar Empleados', level=3)
    doc.add_paragraph('1. Ve al menú y selecciona "Personal > Empleados"')
    doc.add_paragraph('2. Verás una tabla con todos los empleados')
    doc.add_paragraph('3. Puedes filtrar por departamento, sucursal o estado')
    doc.add_paragraph('4. Usa la búsqueda rápida para encontrar un empleado específico')
    
    add_heading_styled(doc, 'Crear Nuevo Empleado', level=3)
    doc.add_paragraph('1. Haz clic en el botón "Nuevo Empleado"')
    doc.add_paragraph('2. Completa los campos requeridos:')
    doc.add_paragraph('Nombre completo', style='List Bullet 2')
    doc.add_paragraph('Email corporativo', style='List Bullet 2')
    doc.add_paragraph('Teléfono de contacto', style='List Bullet 2')
    doc.add_paragraph('Departamento', style='List Bullet 2')
    doc.add_paragraph('Puesto', style='List Bullet 2')
    doc.add_paragraph('Sucursal', style='List Bullet 2')
    doc.add_paragraph('3. Adjunta documentos si es necesario')
    doc.add_paragraph('4. Haz clic en "Guardar"')
    
    add_heading_styled(doc, 'Editar Empleado', level=3)
    doc.add_paragraph('1. Haz clic en el empleado que deseas editar')
    doc.add_paragraph('2. Haz clic en el botón "Editar"')
    doc.add_paragraph('3. Realiza los cambios necesarios')
    doc.add_paragraph('4. Haz clic en "Guardar"')
    
    # Asistencia
    add_heading_styled(doc, '4.3 Asistencia', level=2)
    doc.add_paragraph(
        'Controla la asistencia de tu equipo de forma sencilla.'
    )
    doc.add_paragraph('Opciones disponibles:', style='List Bullet')
    doc.add_paragraph('Ver calendarios de asistencia por empleado', style='List Bullet 2')
    doc.add_paragraph('Registrar asistencia (marcador rápido)', style='List Bullet 2')
    doc.add_paragraph('Ver reportes de asistencia', style='List Bullet 2')
    doc.add_paragraph('Exportar datos de asistencia', style='List Bullet 2')
    
    # Ausencias
    add_heading_styled(doc, '4.4 Ausencias y Vacaciones', level=2)
    doc.add_paragraph('Gestiona los días de vacaciones y ausencias.')
    
    add_heading_styled(doc, 'Solicitar Ausencia', level=3)
    doc.add_paragraph('1. Ve a "Ausencias > Solicitudes"')
    doc.add_paragraph('2. Haz clic en "Nueva Solicitud"')
    doc.add_paragraph('3. Selecciona el tipo de ausencia:')
    doc.add_paragraph('Vacaciones', style='List Bullet 2')
    doc.add_paragraph('Permiso', style='List Bullet 2')
    doc.add_paragraph('Incapacidad', style='List Bullet 2')
    doc.add_paragraph('Otros', style='List Bullet 2')
    doc.add_paragraph('4. Selecciona las fechas')
    doc.add_paragraph('5. Añade un comentario (opcional)')
    doc.add_paragraph('6. Haz clic en "Enviar"')
    
    add_heading_styled(doc, 'Aprobar Ausencias (Gerentes/RRHH)', level=3)
    doc.add_paragraph('1. Ve a "Ausencias > Solicitudes Pendientes"')
    doc.add_paragraph('2. Revisa la solicitud')
    doc.add_paragraph('3. Haz clic en "Aprobar" o "Rechazar"')
    doc.add_paragraph('4. Añade un comentario si lo deseas')
    
    # Objetivos
    add_heading_styled(doc, '4.5 Objetivos y KPIs', level=2)
    doc.add_paragraph('Define y da seguimiento a los objetivos del equipo.')
    
    doc.add_paragraph('Crear Objetivo:', style='List Bullet')
    doc.add_paragraph('1. Ve a "Objetivos"')
    doc.add_paragraph('2. Haz clic en "Nuevo Objetivo"', style='List Bullet 2')
    doc.add_paragraph('3. Define el objetivo, meta y fecha de entrega', style='List Bullet 2')
    doc.add_paragraph('4. Asigna a un empleado', style='List Bullet 2')
    doc.add_paragraph('5. Guarda', style='List Bullet 2')
    
    doc.add_paragraph('Hacer Seguimiento:', style='List Bullet')
    doc.add_paragraph('1. Ve a "Objetivos > En Seguimiento"', style='List Bullet 2')
    doc.add_paragraph('2. Actualiza el progreso', style='List Bullet 2')
    doc.add_paragraph('3. Añade comentarios si es necesario', style='List Bullet 2')
    
    # Tareas
    add_heading_styled(doc, '4.6 Tareas', level=2)
    doc.add_paragraph('Sistema colaborativo de gestión de tareas.')
    
    doc.add_paragraph('Crear Tarea:', style='List Bullet')
    doc.add_paragraph('1. Ve a "Tareas"', style='List Bullet 2')
    doc.add_paragraph('2. Haz clic en "Nueva Tarea"', style='List Bullet 2')
    doc.add_paragraph('3. Nombre de la tarea', style='List Bullet 2')
    doc.add_paragraph('4. Descripción detallada', style='List Bullet 2')
    doc.add_paragraph('5. Asigna a uno o varios usuarios', style='List Bullet 2')
    doc.add_paragraph('6. Define fecha de entrega', style='List Bullet 2')
    doc.add_paragraph('7. Prioridad (Baja, Media, Alta)', style='List Bullet 2')
    
    # Nómina
    add_heading_styled(doc, '4.7 Nómina', level=2)
    doc.add_paragraph('Gestión de nóminas y cálculo de salarios.')
    
    doc.add_paragraph('Consultar Nómina:', style='List Bullet')
    doc.add_paragraph('1. Ve a "Nómina"', style='List Bullet 2')
    doc.add_paragraph('2. Selecciona el mes', style='List Bullet 2')
    doc.add_paragraph('3. Descarga el recibo de nómina en PDF', style='List Bullet 2')
    
    doc.add_paragraph('Generar Nómina (Administrador):', style='List Bullet')
    doc.add_paragraph('1. Ve a "Nómina > Generar"', style='List Bullet 2')
    doc.add_paragraph('2. Selecciona el mes', style='List Bullet 2')
    doc.add_paragraph('3. Revisa los datos', style='List Bullet 2')
    doc.add_paragraph('4. Haz clic en "Procesar"', style='List Bullet 2')
    
    doc.add_page_break()

def add_roles_guide(doc):
    """Add guide by roles"""
    add_heading_styled(doc, '5. Guía por Rol', level=1, color=RGBColor(0, 51, 102))
    
    roles_data = [
        {
            'name': '5.1 Administrador (Dueño de Empresa)',
            'description': 'Tiene control total de la empresa y todos sus datos.',
            'permissions': [
                'Gestionar empleados (crear, editar, eliminar)',
                'Configuración de empresa y sucursales',
                'Gestión de usuarios y permisos',
                'Ver todos los reportes',
                'Generar nóminas',
                'Configurar ausencias y vacaciones',
                'Acceso a auditoría'
            ],
            'common_tasks': [
                'Crear nuevas sucursales',
                'Asignar roles a usuarios',
                'Exportar reportes mensuales',
                'Revisar datos de asistencia',
                'Generar nóminas'
            ]
        },
        {
            'name': '5.2 Recursos Humanos (RRHH)',
            'description': 'Gestiona la operativa diaria de recursos humanos.',
            'permissions': [
                'Gestionar empleados',
                'Revisar y aprobar ausencias',
                'Registrar asistencia',
                'Ver reportes de personal',
                'Gestionar objetivos',
                'Crear tareas',
                'Ver nómina'
            ],
            'common_tasks': [
                'Revisar solicitudes de vacaciones',
                'Registrar ausencias',
                'Crear nuevos empleados',
                'Generar reportes de asistencia',
                'Seguimiento de objetivos'
            ]
        },
        {
            'name': '5.3 Gerente de Equipo',
            'description': 'Supervisa el equipo asignado.',
            'permissions': [
                'Ver empleados de su equipo/sucursal',
                'Ver asistencia de su equipo',
                'Crear y asignar tareas',
                'Definir objetivos para su equipo',
                'Ver nómina de su equipo',
                'Revisar ausencias'
            ],
            'common_tasks': [
                'Asignar tareas a empleados',
                'Ver asistencia del equipo',
                'Crear objetivos mensuales',
                'Revisar progreso de tareas',
                'Validar ausencias'
            ]
        },
        {
            'name': '5.4 Empleado',
            'description': 'Acceso a sus datos personales y funcionalidades básicas.',
            'permissions': [
                'Ver datos personales',
                'Consultar asistencia propia',
                'Solicitar vacaciones/ausencias',
                'Ver tareas asignadas',
                'Consultar su nómina',
                'Actualizar perfil'
            ],
            'common_tasks': [
                'Solicitar vacaciones',
                'Ver tareas pendientes',
                'Actualizar estado de tareas',
                'Consultar asistencia',
                'Descargar recibo de nómina'
            ]
        }
    ]
    
    for role in roles_data:
        add_heading_styled(doc, role['name'], level=2)
        doc.add_paragraph(role['description'])
        
        doc.add_paragraph('Permisos:', style='List Bullet')
        for perm in role['permissions']:
            doc.add_paragraph(perm, style='List Bullet 2')
        
        doc.add_paragraph('Tareas Comunes:', style='List Bullet')
        for task in role['common_tasks']:
            doc.add_paragraph(task, style='List Bullet 2')
        
        doc.add_paragraph('')
    
    doc.add_page_break()

def add_common_operations(doc):
    """Add common operations section"""
    add_heading_styled(doc, '6. Operaciones Comunes', level=1, color=RGBColor(0, 51, 102))
    
    add_heading_styled(doc, 'Búsqueda y Filtros', level=2)
    doc.add_paragraph('Casi todos los listados tienen opciones de búsqueda y filtro:')
    doc.add_paragraph('1. Usa la barra de búsqueda en la parte superior para buscar por nombre, email, etc.', style='List Number')
    doc.add_paragraph('2. Usa los filtros para refinar resultados por departamento, estado, etc.', style='List Number')
    doc.add_paragraph('3. Los resultados se actualizan automáticamente', style='List Number')
    
    add_heading_styled(doc, 'Exportar Datos', level=2)
    doc.add_paragraph('Para exportar datos a Excel:')
    doc.add_paragraph('1. Ve al listado que deseas exportar', style='List Number')
    doc.add_paragraph('2. Haz clic en el botón "Exportar" (ícono de descarga)', style='List Number')
    doc.add_paragraph('3. Selecciona formato: Excel, PDF o CSV', style='List Number')
    doc.add_paragraph('4. El archivo se descargará automáticamente', style='List Number')
    
    add_heading_styled(doc, 'Importar Datos', level=2)
    doc.add_paragraph('Para importar empleados en lote:')
    doc.add_paragraph('1. Ve a "Personal > Importar Empleados"', style='List Number')
    doc.add_paragraph('2. Descarga la plantilla de ejemplo', style='List Number')
    doc.add_paragraph('3. Completa los datos en la plantilla', style='List Number')
    doc.add_paragraph('4. Sube el archivo', style='List Number')
    doc.add_paragraph('5. Revisa los datos antes de confirmar', style='List Number')
    
    add_heading_styled(doc, 'Cambiar Contraseña', level=2)
    doc.add_paragraph('1. Haz clic en tu foto de perfil (esquina superior derecha)', style='List Number')
    doc.add_paragraph('2. Selecciona "Mi Perfil"', style='List Number')
    doc.add_paragraph('3. Haz clic en "Cambiar Contraseña"', style='List Number')
    doc.add_paragraph('4. Ingresa tu contraseña actual', style='List Number')
    doc.add_paragraph('5. Ingresa la nueva contraseña dos veces', style='List Number')
    doc.add_paragraph('6. Haz clic en "Actualizar"', style='List Number')
    
    add_heading_styled(doc, 'Cerrar Sesión', level=2)
    doc.add_paragraph('1. Haz clic en tu foto de perfil (esquina superior derecha)', style='List Number')
    doc.add_paragraph('2. Haz clic en "Cerrar Sesión"', style='List Number')
    
    doc.add_page_break()

def add_faq(doc):
    """Add FAQ section"""
    add_heading_styled(doc, '7. Preguntas Frecuentes', level=1, color=RGBColor(0, 51, 102))
    
    faqs = [
        {
            'q': '¿Cómo agrego un nuevo empleado?',
            'a': 'Ve a Personal > Empleados, haz clic en "Nuevo Empleado" y completa el formulario con los datos solicitados.'
        },
        {
            'q': '¿Puedo editar información de un empleado?',
            'a': 'Sí, ve a Personal > Empleados, haz clic en el empleado y luego en "Editar". Realiza los cambios y guarda.'
        },
        {
            'q': '¿Cómo solicito una ausencia o vacaciones?',
            'a': 'Ve a Ausencias > Mis Solicitudes, haz clic en "Nueva Solicitud", selecciona el tipo y fechas, luego envía.'
        },
        {
            'q': '¿Dónde veo mis tareas asignadas?',
            'a': 'Ve al módulo Tareas. Verás todas las tareas asignadas a ti con su estado y fecha de entrega.'
        },
        {
            'q': '¿Cómo descargo mi recibo de nómina?',
            'a': 'Ve a Nómina, selecciona el mes deseado y haz clic en "Descargar PDF".'
        },
        {
            'q': '¿Qué hago si tengo problemas con mi contraseña?',
            'a': 'En la pantalla de login, haz clic en "¿Olvidaste tu contraseña?" y sigue las instrucciones enviadas al email.'
        },
        {
            'q': '¿Puedo cambiar mi foto de perfil?',
            'a': 'Sí, ve a Mi Perfil, haz clic en tu foto actual y sube una nueva imagen.'
        },
        {
            'q': '¿Cómo genero un reporte?',
            'a': 'Ve al módulo que deseas reportar, selecciona los filtros deseados y haz clic en "Exportar".'
        },
        {
            'q': '¿Quién puede aprobar mis ausencias?',
            'a': 'Tu gerente o alguien del área de RRHH. Depende de la configuración de tu empresa.'
        },
        {
            'q': '¿Qué hago si veo un error?',
            'a': 'Intenta recargar la página. Si persiste, contacta a tu administrador o equipo de soporte.'
        }
    ]
    
    for idx, faq in enumerate(faqs, 1):
        p = doc.add_paragraph()
        p.add_run(f'P: {faq["q"]}').bold = True
        doc.add_paragraph(f'R: {faq["a"]}', style='List Bullet')
        doc.add_paragraph('')

def add_troubleshooting(doc):
    """Add troubleshooting section"""
    doc.add_page_break()
    add_heading_styled(doc, '8. Solución de Problemas', level=1, color=RGBColor(0, 51, 102))
    
    problems = [
        {
            'problem': 'No puedo iniciar sesión',
            'solutions': [
                'Verifica que tu usuario/email y contraseña sean correctos',
                'Asegúrate de que Caps Lock no está activo',
                'Intenta usar tu email en lugar de tu usuario',
                'Si olvidaste la contraseña, usa la opción "¿Olvidaste tu contraseña?"',
                'Contacta a tu administrador si aún así no funciona'
            ]
        },
        {
            'problem': 'Veo un error al guardar un formulario',
            'solutions': [
                'Revisa que todos los campos obligatorios estén completados',
                'Verifica que el formato de los datos sea correcto (ej: emails válidos)',
                'Intenta recargar la página',
                'Borra el caché del navegador',
                'Intenta desde otro navegador'
            ]
        },
        {
            'problem': 'No veo algunos módulos en el menú',
            'solutions': [
                'Esto depende de tu rol y permisos',
                'Contacta a tu administrador para solicitar acceso',
                'Verifica que tu rol tenga los permisos necesarios'
            ]
        },
        {
            'problem': 'Los datos no se actualizan',
            'solutions': [
                'Intenta recargar la página (F5)',
                'Cierra la sesión y vuelve a iniciarla',
                'Borra el caché del navegador',
                'Intenta desde otro dispositivo'
            ]
        },
        {
            'problem': 'No puedo descargar reportes',
            'solutions': [
                'Verifica que tengas permisos para exportar datos',
                'Intenta desde otro navegador',
                'Desactiva las extensiones del navegador que bloqueen descargas',
                'Verifica que tu antivirus no bloquee descargas'
            ]
        },
        {
            'problem': 'El sitio está lento',
            'solutions': [
                'Intenta cerrar otras pestañas del navegador',
                'Recarga la página',
                'Verifica tu conexión a internet',
                'Intenta desde otro dispositivo',
                'Contacta a tu administrador'
            ]
        }
    ]
    
    for problem in problems:
        p = doc.add_paragraph()
        p.add_run(f'Problema: {problem["problem"]}').bold = True
        
        doc.add_paragraph('Soluciones:', style='List Bullet')
        for solution in problem['solutions']:
            doc.add_paragraph(solution, style='List Bullet 2')
        
        doc.add_paragraph('')

def generate_manual():
    """Main function to generate the manual"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    print("Generando manual de usuario...")
    
    # Add all sections
    add_title_page(doc)
    print("✓ Página de título añadida")
    
    add_table_of_contents(doc)
    print("✓ Tabla de contenidos añadida")
    
    add_introduction(doc)
    print("✓ Introducción añadida")
    
    add_login_section(doc)
    print("✓ Sección de acceso añadida")
    
    add_main_interface(doc)
    print("✓ Interfaz principal añadida")
    
    add_modules_section(doc)
    print("✓ Módulos principales añadidos")
    
    add_roles_guide(doc)
    print("✓ Guía por roles añadida")
    
    add_common_operations(doc)
    print("✓ Operaciones comunes añadidas")
    
    add_faq(doc)
    print("✓ Preguntas frecuentes añadidas")
    
    add_troubleshooting(doc)
    print("✓ Solución de problemas añadida")
    
    # Save document
    output_path = 'Manual_Usuario_PuntoPymes.docx'
    doc.save(output_path)
    print(f"\n✓ Manual generado exitosamente: {output_path}")
    print(f"✓ El archivo está en: {output_path}")

if __name__ == '__main__':
    generate_manual()
